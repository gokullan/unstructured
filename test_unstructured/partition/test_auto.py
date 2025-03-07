import os
import pathlib
import pytest

import docx

from unstructured.documents.elements import NarrativeText, Title, Text, ListItem
from unstructured.partition.auto import partition
import unstructured.partition.auto as auto

EXAMPLE_DOCS_DIRECTORY = pathlib.Path(__file__).parent.resolve()


EXPECTED_EMAIL_OUTPUT = [
    NarrativeText(text="This is a test email to use for unit tests."),
    Title(text="Important points:"),
    ListItem(text="Roses are red"),
    ListItem(text="Violets are blue"),
]


def test_auto_partition_email_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-email.eml")
    elements = partition(filename=filename)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT


def test_auto_partition_email_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-email.eml")
    with open(filename, "r") as f:
        elements = partition(file=f)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT


def test_auto_partition_email_from_file_rb():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-email.eml")
    with open(filename, "rb") as f:
        elements = partition(file=f)
    assert len(elements) > 0
    assert elements == EXPECTED_EMAIL_OUTPUT


@pytest.fixture
def mock_docx_document():
    document = docx.Document()

    document.add_paragraph("These are a few of my favorite things:", style="Heading 1")
    # NOTE(robinson) - this should get picked up as a list item due to the •
    document.add_paragraph("• Parrots", style="Normal")
    document.add_paragraph("Hockey", style="List Bullet")
    # NOTE(robinson) - this should get picked up as a title
    document.add_paragraph("Analysis", style="Normal")
    # NOTE(robinson) - this should get dropped because it is empty
    document.add_paragraph("", style="Normal")
    # NOTE(robinson) - this should get picked up as a narrative text
    document.add_paragraph("This is my first thought. This is my second thought.", style="Normal")
    document.add_paragraph("This is my third thought.", style="Body Text")
    # NOTE(robinson) - this should just be regular text
    document.add_paragraph("2023")

    return document


@pytest.fixture
def expected_docx_elements():
    return [
        Title("These are a few of my favorite things:"),
        ListItem("Parrots"),
        ListItem("Hockey"),
        Title("Analysis"),
        NarrativeText("This is my first thought. This is my second thought."),
        NarrativeText("This is my third thought."),
        Text("2023"),
    ]


def test_auto_partition_docx_with_filename(mock_docx_document, expected_docx_elements, tmpdir):
    filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    mock_docx_document.save(filename)

    elements = partition(filename=filename)
    assert elements == expected_docx_elements


def test_auto_partition_docx_with_file(mock_docx_document, expected_docx_elements, tmpdir):
    filename = os.path.join(tmpdir.dirname, "mock_document.docx")
    mock_docx_document.save(filename)

    with open(filename, "rb") as f:
        elements = partition(file=f)
    assert elements == expected_docx_elements


def test_auto_partition_html_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "example-10k.html")
    elements = partition(filename=filename)
    assert len(elements) > 0


def test_auto_partition_html_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-html.html")
    with open(filename, "r") as f:
        elements = partition(file=f)
    assert len(elements) > 0


def test_auto_partition_html_from_file_rb():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-html.html")
    with open(filename, "rb") as f:
        elements = partition(file=f)
    assert len(elements) > 0


EXPECTED_TEXT_OUTPUT = [
    NarrativeText(text="This is a test document to use for unit tests."),
    Title(text="Important points:"),
    ListItem(text="Hamburgers are delicious"),
    ListItem(text="Dogs are the best"),
    ListItem(text="I love fuzzy blankets"),
]


def test_auto_partition_text_from_filename():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-text.txt")
    elements = partition(filename=filename)
    assert len(elements) > 0
    assert elements == EXPECTED_TEXT_OUTPUT


def test_auto_partition_text_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "fake-text.txt")
    with open(filename, "r") as f:
        elements = partition(file=f)
    assert len(elements) > 0
    assert elements == EXPECTED_TEXT_OUTPUT
    filename = os.path.join(
        EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "layout-parser-paper-fast.pdf"
    )
    elements = partition(filename=filename)
    assert len(elements) > 0


def test_auto_partition_pdf_from_file():
    filename = os.path.join(
        EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "layout-parser-paper-fast.pdf"
    )
    with open(filename, "rb") as f:
        elements = partition(file=f)
    assert len(elements) > 0


def test_auto_partition_jpg():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "example.jpg")
    elements = partition(filename=filename)
    assert len(elements) > 0


def test_auto_partition_jpg_from_file():
    filename = os.path.join(EXAMPLE_DOCS_DIRECTORY, "..", "..", "example-docs", "example.jpg")
    with open(filename, "rb") as f:
        elements = partition(file=f)
    assert len(elements) > 0


def test_auto_partition_raises_with_bad_type(monkeypatch):
    monkeypatch.setattr(auto, "detect_filetype", lambda *args, **kwargs: None)
    with pytest.raises(ValueError):
        partition(filename="made-up.fake")
